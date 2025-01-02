import os
from flask import Flask, send_file
import time
import openai
import csv
from docx import Document
from docx.shared import Pt
from typing import List, Dict

app = Flask(__name__)

@app.route("/")
def index():
    return send_file('src/index.html')

class ChatGPTQuestionAutomation:
    def __init__(self, api_key: str, model: str = "gpt-3.5-turbo"):
        """
        Initialize the ChatGPT Question Automation script
        
        :param api_key: OpenAI API key
        :param model: OpenAI model to use (defaults to latest GPT-4 Turbo)
        """
        openai.api_key = 'YOUR API KEY'
        self.model = model
        
        # Validate model selection
        valid_models = [
            "gpt-3.5-turbo", 
            "gpt-4", 
            "gpt-3.5-turbo", 
            "gpt-4-1106-preview"
        ]
        if self.model not in valid_models:
            raise ValueError(f"Invalid model. Choose from: {', '.join(valid_models)}")
    
    def load_questions_from_csv(self, filename: str) -> List[Dict]:
        """
        Load questions from a CSV file
        
        :param filename: Path to the CSV file
        :return: List of question dictionaries
        """
        questions = []
        try:
            with open(filename, 'r', encoding='utf-8') as file:
                reader = csv.DictReader(file)
                for row in reader:
                    questions.append({
                        'category': row['category'],
                        'question': row['question']
                    })
            print(f"Successfully loaded {len(questions)} questions from {filename}")
            return questions
        except FileNotFoundError:
            print(f"Error: File {filename} not found")
            return []
        except Exception as e:
            print(f"Error loading questions: {e}")
            return []
    
    def ask_chatgpt(self, question: str) -> str:
        """
        Send a question to ChatGPT and retrieve the response
        
        :param question: Question to ask
        :return: ChatGPT's response
        """
        try:
            response = openai.ChatCompletion.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are a helpful assistant."},
                    {"role": "user", "content": question}
                ],
                max_tokens=500,  # Adjust as needed
                temperature=0.7
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            print(f"Error asking question: {e}")
            return f"Error: {str(e)}"
    
    def save_to_word(self, questions_and_responses: List[Dict], output_file: str = "chatgpt_responses.docx"):
        """
        Save questions and responses to a Word document
        
        :param questions_and_responses: List of dictionaries containing questions and responses
        :param output_file: Name of the output Word document
        """
        document = Document()
        
        # Set default font and size
        style = document.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Group questions by category
        categories = {}
        for entry in questions_and_responses:
            category = entry['category']
            if category not in categories:
                categories[category] = []
            categories[category].append(entry)
        
        # Add content by category
        for category, entries in categories.items():
            # Add category header
            document.add_heading(category, level=1)
            
            for entry in entries:
                # Add question
                document.add_paragraph(f"Question: {entry['question']}", style='Quote')
                
                # Add response
                document.add_paragraph(f"Solution: {entry['response']}")
                
                # Add spacing between Q&A pairs
                document.add_paragraph()
            
            # Add page break after each category
            document.add_page_break()
        
        document.save(output_file)
        print(f"Responses saved to {output_file}")
    
    def run_automation(self, questions_file: str):
        """
        Main automation method to load questions, get responses, and save to document
        
        :param questions_file: Path to the CSV file containing questions
        """
        print(f"Starting automation with model: {self.model}")
        
        # Load questions from CSV
        questions = self.load_questions_from_csv(questions_file)
        if not questions:
            print("No questions loaded. Exiting.")
            return
        
        # Collect responses
        responses = []
        total_questions = len(questions)
        
        for i, q in enumerate(questions, 1):
            print(f"Processing question {i} of {total_questions}")
            response = self.ask_chatgpt(q['question'])
            responses.append({
                'category': q['category'],
                'question': q['question'],
                'response': response
            })
            
            # Rate limiting to avoid overwhelming the API
            time.sleep(2)
        
        # Save to Word document
        self.save_to_word(responses)

def main():
    # Your API key should be in quotes as it's a string
    API_KEY = "YOUR API KEY"
    
    # Path to your questions CSV file
    QUESTIONS_FILE = 'questions.csv'
    
    # Initialize and run automation
    automation = ChatGPTQuestionAutomation(api_key=API_KEY)
    automation.run_automation(QUESTIONS_FILE)

if __name__ == "__main__":
    main()